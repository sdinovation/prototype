"""
RAG（检索增强生成）智能问答系统
使用 Flask + Sentence Transformers + FAISS 构建
"""

import os
import numpy as np
from flask import Flask, render_template, request, jsonify
from sentence_transformers import SentenceTransformer
import faiss

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局变量 =====================
model = None
index = None
documents = None
EMBEDDING_DIM = 384  # all-MiniLM-L6-v2 的维度


# ===================== 加载知识库 =====================
def load_knowledge_base(filepath='knowledge_base.txt'):
    """从文本文件加载知识库，按段落分割"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    # 按空行分割成段落
    docs = [p.strip() for p in content.split('\n\n') if p.strip()]
    return docs


# ===================== 构建 FAISS 索引 =====================
def build_faiss_index(documents, model):
    """将文档转换为向量并构建 FAISS 索引"""
    print("正在生成文档向量...")
    embeddings = model.encode(documents, show_progress_bar=True)
    # 归一化向量（用于余弦相似度）
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    # 构建 FAISS 索引（使用内积索引，等价于归一化后的余弦相似度）
    dim = embeddings.shape[1]
    idx = faiss.IndexFlatIP(dim)
    idx.add(embeddings)
    return idx, embeddings


# ===================== 检索相关文档 =====================
def retrieve(query, model, index, documents, top_k=3):
    """根据查询检索最相关的文档"""
    # 将查询转换为向量
    query_embedding = model.encode([query])
    query_embedding = np.array(query_embedding, dtype=np.float32)
    faiss.normalize_L2(query_embedding)
    # 在 FAISS 索引中搜索
    scores, indices = index.search(query_embedding, top_k)
    results = []
    for i in range(len(indices[0])):
        idx = indices[0][i]
        score = float(scores[0][i])
        if idx >= 0 and idx < len(documents):
            results.append({
                'text': documents[idx],
                'score': round(score, 4)
            })
    return results


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs):
    """基于检索到的文档生成回答（简易模板方式，不依赖大语言模型）"""
    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。"
    
    # 构建上下文
    context = "\n".join([f"[参考资料{i+1}] {doc['text']}" for i, doc in enumerate(retrieved_docs)])
    
    # 简易回答生成（基于模板拼接）
    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.2%}）\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，从本地知识库中检索最相关的文档片段。"
    
    return answer


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    
    if not question:
        return jsonify({'error': '请输入问题'}), 400
    
    # 检索相关文档
    results = retrieve(question, model, index, documents, top_k=3)
    
    # 生成回答
    answer = generate_answer(question, results)
    
    return jsonify({
        'question': question,
        'answer': answer,
        'sources': results
    })


@app.route('/api/status')
def status():
    """返回系统状态"""
    return jsonify({
        'status': 'running',
        'document_count': len(documents),
        'model_name': model.__class__.__name__
    })


# ===================== 初始化 =====================
if __name__ == '__main__':
    print("=" * 50)
    print("  RAG 检索增强生成系统")
    print("  Flask + Sentence Transformers + FAISS")
    print("=" * 50)
    
    # 加载 Sentence Transformer 模型
    print("\n[1/3] 正在加载 Sentence Transformer 模型...")
    model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
    print(f"  模型加载完成，向量维度: {model.get_sentence_embedding_dimension()}")
    
    # 加载知识库
    print("\n[2/3] 正在加载知识库...")
    documents = load_knowledge_base()
    print(f"  知识库加载完成，共 {len(documents)} 条文档")
    
    # 构建 FAISS 索引
    print("\n[3/3] 正在构建 FAISS 向量索引...")
    index, embeddings = build_faiss_index(documents, model)
    print(f"  索引构建完成，包含 {index.ntotal} 条向量")
    
    print("\n" + "=" * 50)
    print("  系统启动成功！请访问 http://localhost:5000")
    print("=" * 50 + "\n")
    
    # 启动 Flask 服务
    app.run(host='0.0.0.0', port=5000, debug=False)
"""
RAG（检索增强生成）智能问答"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfpl"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text,"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n',"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  #"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df':"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0."""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score ="""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) +"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1],"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k *"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices["""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results ="""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x["""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices["""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score /"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ================="""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到',"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让',"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个',"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f""""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer +="""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n"."""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ================="""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error':"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer,"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question,"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['."""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共解析 {len(chunks)} 个文本片段',
        'doc_id': doc_record['id'],
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks'])"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共解析 {len(chunks)} 个文本片段',
        'doc_id': doc_record['id'],
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks']) for d in doc_list)
    })


@app.route"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共解析 {len(chunks)} 个文本片段',
        'doc_id': doc_record['id'],
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks']) for d in doc_list)
    })


@app.route('/api/docs', methods=['GET'])
def get_docs():
    """获取知识库文档列表"""
    result = []
    for doc in doc_list:
        result.append({
            'id': doc"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共解析 {len(chunks)} 个文本片段',
        'doc_id': doc_record['id'],
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks']) for d in doc_list)
    })


@app.route('/api/docs', methods=['GET'])
def get_docs():
    """获取知识库文档列表"""
    result = []
    for doc in doc_list:
        result.append({
            'id': doc['id'],
            'source': doc['source'],
            'chunk_count': len(doc['chunks']),
            'upload_time': doc.get('upload"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共解析 {len(chunks)} 个文本片段',
        'doc_id': doc_record['id'],
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks']) for d in doc_list)
    })


@app.route('/api/docs', methods=['GET'])
def get_docs():
    """获取知识库文档列表"""
    result = []
    for doc in doc_list:
        result.append({
            'id': doc['id'],
            'source': doc['source'],
            'chunk_count': len(doc['chunks']),
            'upload_time': doc.get('upload_time', '')
        })
    return jsonify({'docs': result, 'total': len(result)})


@app.route('/api/docs/<doc_id>', methods=['DELETE'])
def delete_doc(doc_id):
    """删除知识库中的文档"""
    global doc"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共解析 {len(chunks)} 个文本片段',
        'doc_id': doc_record['id'],
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks']) for d in doc_list)
    })


@app.route('/api/docs', methods=['GET'])
def get_docs():
    """获取知识库文档列表"""
    result = []
    for doc in doc_list:
        result.append({
            'id': doc['id'],
            'source': doc['source'],
            'chunk_count': len(doc['chunks']),
            'upload_time': doc.get('upload_time', '')
        })
    return jsonify({'docs': result, 'total': len(result)})


@app.route('/api/docs/<doc_id>', methods=['DELETE'])
def delete_doc(doc_id):
    """删除知识库中的文档"""
    global doc_list
    original_len = len(doc_list)
    doc_list = [d for d in doc_list if d['id'] != doc_id]
    if len(doc_list) == original_len:
        return jsonify({'error': '文档不存在'}), 404
    rebuild_index()
"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共解析 {len(chunks)} 个文本片段',
        'doc_id': doc_record['id'],
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks']) for d in doc_list)
    })


@app.route('/api/docs', methods=['GET'])
def get_docs():
    """获取知识库文档列表"""
    result = []
    for doc in doc_list:
        result.append({
            'id': doc['id'],
            'source': doc['source'],
            'chunk_count': len(doc['chunks']),
            'upload_time': doc.get('upload_time', '')
        })
    return jsonify({'docs': result, 'total': len(result)})


@app.route('/api/docs/<doc_id>', methods=['DELETE'])
def delete_doc(doc_id):
    """删除知识库中的文档"""
    global doc_list
    original_len = len(doc_list)
    doc_list = [d for d in doc_list if d['id'] != doc_id]
    if len(doc_list) == original_len:
        return jsonify({'error': '文档不存在'}), 404
    rebuild_index()
    return jsonify({
        'success': True,
        'message': '文档已删除',
        'total_docs': len(doc_list),
        'total"""
RAG（检索增强生成）智能问答系统 - 完整版
功能：多格式文档上传、BM25混合检索、关键词高亮、知识库管理、
      对话模式切换、导出对话记录
技术栈：Flask + Sentence Transformers + FAISS + jieba(BM25)
"""

import os
import re
import math
import uuid
import time
import json
from collections import defaultdict
from datetime import datetime

import numpy as np
from flask import Flask, render_template, request, jsonify, send_file
from sentence_transformers import SentenceTransformer
import faiss
import jieba

# 设置 HuggingFace 镜像
os.environ['HF_ENDPOINT'] = 'https://hf-mirror.com'

app = Flask(__name__)

# ===================== 全局状态 =====================
model = None
faiss_index = None
doc_list = []          # 所有文档片段 [{"id":, "text":, "source":, "chunks": []}]
bm25_index = None      # BM25倒排索引
doc_tokens_list = []   # BM25用的分词结果
chat_history = []       # 对话历史 [{"id":, "mode":, "question":, "answer":, "time":}]
UPLOAD_FOLDER = 'uploads'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


# ===================== 文档解析 =====================
def parse_file(filepath):
    """根据文件类型解析文档内容，返回文本字符串"""
    ext = os.path.splitext(filepath)[1].lower()
    if ext in ['.txt', '.md']:
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()
    elif ext == '.pdf':
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(filepath)
            text = ""
            for page in doc:
                text += page.get_text()
            doc.close()
            return text
        except ImportError:
            # 备用方案：pdfplumber
            try:
                import pdfplumber
                text = ""
                with pdfplumber.open(filepath) as pdf:
                    for page in pdf.pages:
                        text += page.extract_text() or ""
                return text
            except ImportError:
                return f"[PDF解析失败] 请安装 PyMuPDF: pip install PyMuPDF"
    elif ext == '.docx':
        try:
            from docx import Document
            doc = Document(filepath)
            return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except ImportError:
            return "[DOCX解析失败] 请安装 python-docx: pip install python-docx"
    else:
        return f"[不支持的文件格式] {ext}"


def split_documents(text, source_name="unknown"):
    """将文本智能分割成文档片段"""
    # 按空行或多个换行分割
    raw_chunks = re.split(r'\n\s*\n', text)
    chunks = []
    for chunk in raw_chunks:
        chunk = chunk.strip()
        if len(chunk) < 5:  # 过滤太短的片段
            continue
        # 如果单段过长（>500字），按句子进一步分割
        if len(chunk) > 500:
            sentences = re.split(r'(?<=[。！？\.\!\?])', chunk)
            sub = ""
            for s in sentences:
                sub += s
                if len(sub) > 300:
                    chunks.append(sub.strip())
                    sub = ""
            if sub.strip():
                chunks.append(sub.strip())
        else:
            chunks.append(chunk)
    return chunks


# ===================== 知识库管理 =====================
def rebuild_index():
    """重建FAISS索引和BM25索引"""
    global faiss_index, bm25_index, doc_tokens_list
    if not doc_list:
        faiss_index = None
        bm25_index = None
        doc_tokens_list = []
        return
    all_texts = []
    doc_tokens_list = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_texts.append(chunk)
            tokens = list(jieba.cut(chunk))
            doc_tokens_list.append(tokens)
    # FAISS向量索引
    embeddings = model.encode(all_texts, show_progress_bar=False)
    embeddings = np.array(embeddings, dtype=np.float32)
    faiss.normalize_L2(embeddings)
    dim = embeddings.shape[1]
    faiss_index = faiss.IndexFlatIP(dim)
    faiss_index.add(embeddings)
    # BM25索引
    bm25_index = build_bm25(doc_tokens_list)


# ===================== BM25 检索 =====================
def build_bm25(doc_tokens_list):
    """构建BM25倒排索引"""
    N = len(doc_tokens_list)  # 文档总数
    df = defaultdict(int)     # 词频文档数
    doc_len = []              # 每篇文档长度
    for tokens in doc_tokens_list:
        doc_len.append(len(tokens))
        seen = set()
        for t in tokens:
            if t not in seen:
                df[t] += 1
                seen.add(t)
    avgdl = sum(doc_len) / N if N > 0 else 1
    return {
        'N': N, 'df': dict(df), 'doc_len': doc_len, 'avgdl': avgdl,
        'doc_tokens': doc_tokens_list
    }


def bm25_search(bm25_idx, query_tokens, top_k=10):
    """BM25检索"""
    if not bm25_idx or bm25_idx['N'] == 0:
        return []
    scores = []
    k1 = 1.5
    b = 0.75
    N = bm25_idx['N']
    df = bm25_idx['df']
    avgdl = bm25_idx['avgdl']
    for i, doc_tokens in enumerate(bm25_idx['doc_tokens']):
        score = 0.0
        tf_map = defaultdict(int)
        for t in doc_tokens:
            tf_map[t] += 1
        dl = bm25_idx['doc_len'][i]
        for qt in query_tokens:
            if qt not in tf_map:
                continue
            tf = tf_map[qt]
            idf = math.log((N - df.get(qt, 0) + 0.5) / (df.get(qt, 0) + 0.5) + 1)
            numerator = tf * (k1 + 1)
            denominator = tf + k1 * (1 - b + b * dl / avgdl)
            score += idf * numerator / denominator
        scores.append((i, score))
    scores.sort(key=lambda x: x[1], reverse=True)
    return scores[:top_k]


# ===================== 混合检索（向量 + BM25）=====================
def hybrid_retrieve(query, top_k=3):
    """混合检索：FAISS向量检索 + BM25关键词检索 + RRF融合"""
    all_chunks = []
    for doc in doc_list:
        for chunk in doc['chunks']:
            all_chunks.append(chunk)

    # ---- 向量检索 ----
    query_emb = model.encode([query])
    query_emb = np.array(query_emb, dtype=np.float32)
    faiss.normalize_L2(query_emb)
    vector_scores, vector_indices = faiss_index.search(query_emb, min(top_k * 3, len(all_chunks)))
    vector_results = {}
    for i in range(len(vector_indices[0])):
        idx = vector_indices[0][i]
        score = float(vector_scores[0][i])
        if 0 <= idx < len(all_chunks):
            vector_results[idx] = score

    # ---- BM25检索 ----
    query_tokens = list(jieba.cut(query))
    bm25_raw = bm25_search(bm25_index, query_tokens, top_k=top_k * 3)
    bm25_results = {}
    for idx, score in bm25_raw:
        bm25_results[idx] = score

    # ---- RRF融合（Reciprocal Rank Fusion）----
    all_indices = set(vector_results.keys()) | set(bm25_results.keys())
    rrf_scores = {}
    for idx in all_indices:
        rrf = 0.0
        # 向量排名
        if idx in vector_results:
            rank = sorted(vector_results.items(), key=lambda x: -x[1]).index((idx, vector_results[idx])) if (idx, vector_results[idx]) in sorted(vector_results.items(), key=lambda x: -x[1]) else len(vector_results)
            rrf += 1.0 / (rank + 60)
        # BM25排名
        if idx in bm25_results:
            rank = sorted(bm25_results.items(), key=lambda x: -x[1]).index((idx, bm25_results[idx])) if (idx, bm25_results[idx]) in sorted(bm25_results.items(), key=lambda x: -x[1]) else len(bm25_results)
            rrf += 1.0 / (rank + 60)
        rrf_scores[idx] = rrf

    # 排序取top_k
    sorted_indices = sorted(rrf_scores.items(), key=lambda x: -x[1])[:top_k]

    results = []
    for idx, rrf_score in sorted_indices:
        chunk_text = all_chunks[idx]
        # 查找来源
        source = ""
        for doc in doc_list:
            if chunk_text in doc['chunks']:
                source = doc['source']
                break
        # 归一化相似度显示（RRF分数转百分比）
        max_rrf = sorted_indices[0][1] if sorted_indices else 1
        display_score = round(rrf_score / max_rrf, 4) if max_rrf > 0 else 0
        results.append({
            'text': chunk_text,
            'score': display_score,
            'source': source
        })
    return results


# ===================== 关键词高亮 =====================
def highlight_keywords(text, query):
    """在文本中高亮查询关键词"""
    query_tokens = set(jieba.cut(query))
    # 过滤停用词和单字
    stop_words = {'的', '了', '是', '在', '有', '和', '就', '不', '人', '都', '一', '一个',
                  '上', '也', '很', '到', '说', '要', '去', '你', '会', '着', '没有', '看',
                  '好', '自己', '这', '他', '她', '它', '我', '什么', '吗', '吧', '呢', '啊',
                  '怎么', '哪', '那', '被', '把', '让', '给', '从', '对', '但', '又', '还',
                  '而', '且', '或', '与', '能', '可以', '这个', '那个', '为', '以'}
    keywords = {t for t in query_tokens if len(t) > 1 and t not in stop_words}
    if not keywords:
        return text
    # 对每个关键词用<em>标签包裹
    for kw in keywords:
        text = re.sub(
            re.escape(kw),
            f'<em class="hl">{kw}</em>',
            text
        )
    return text


# ===================== 生成回答 =====================
def generate_answer(query, retrieved_docs, mode='rag'):
    """生成回答"""
    if mode == 'chat':
        return f"你好！我是RAG智能助手（当前为自由聊天模式）。\n\n你的问题是："{query}"\n\n在自由聊天模式下，我暂时无法连接大语言模型API，但我可以告诉你：切换到"知识库问答"模式，我就能从知识库中为你检索专业的答案！"

    if not retrieved_docs:
        return "抱歉，在知识库中没有找到与您的问题相关的信息。您可以尝试上传更多文档到知识库，或者换一种方式描述您的问题。"

    answer = f"根据知识库检索，为您找到以下相关信息：\n\n"
    for i, doc in enumerate(retrieved_docs):
        source_info = f"（来源：{doc['source']}）" if doc.get('source') else ""
        answer += f"【相关文档 {i+1}】（相似度：{doc['score']:.0%}）{source_info}\n{doc['text']}\n\n"
    answer += "---\n"
    answer += "以上回答基于RAG检索增强生成技术，采用向量检索+BM25混合检索策略。"
    return answer


# ===================== 导出对话记录 =====================
def export_chat_history(format_type='markdown'):
    """导出对话记录为Markdown"""
    if not chat_history:
        return None
    lines = []
    lines.append(f"# RAG智能问答系统 - 对话记录")
    lines.append(f"\n导出时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append(f"对话总数：{len(chat_history)} 条\n")
    lines.append("---\n")
    for item in chat_history:
        mode_label = "知识库问答" if item['mode'] == 'rag' else "自由聊天"
        lines.append(f"### [{mode_label}] {item['question']}")
        lines.append(f"\n> {item['time']}\n")
        lines.append(f"{item['answer']}\n")
        lines.append("---\n")
    content = "\n".join(lines)
    return content


# ===================== 加载初始知识库 =====================
def load_initial_knowledge_base():
    """加载默认知识库文件"""
    filepath = 'knowledge_base.txt'
    if not os.path.exists(filepath):
        return
    with open(filepath, 'r', encoding='utf-8') as f:
        text = f.read()
    chunks = split_documents(text, '默认知识库.txt')
    if chunks:
        doc_list.append({
            'id': str(uuid.uuid4())[:8],
            'source': '默认知识库.txt',
            'chunks': chunks,
            'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        })


# ===================== Flask 路由 =====================
@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/api/query', methods=['POST'])
def query():
    """处理用户查询"""
    data = request.get_json()
    question = data.get('question', '').strip()
    mode = data.get('mode', 'rag')  # rag 或 chat

    if not question:
        return jsonify({'error': '请输入问题'}), 400

    if mode == 'chat':
        answer = generate_answer(question, [], mode='chat')
    else:
        if not doc_list:
            answer = "知识库为空，请先上传文档或使用默认知识库。"
        else:
            results = hybrid_retrieve(question, top_k=3)
            answer = generate_answer(question, results, mode='rag')

    # 高亮关键词
    highlighted_answer = highlight_keywords(answer, question)

    # 保存对话历史
    record = {
        'id': str(uuid.uuid4())[:8],
        'mode': mode,
        'question': question,
        'answer': answer,
        'highlighted': highlighted_answer,
        'time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    chat_history.append(record)

    # 获取检索结果
    sources = []
    if mode == 'rag' and doc_list:
        sources = hybrid_retrieve(question, top_k=3)

    return jsonify({
        'question': question,
        'answer': highlighted_answer,
        'mode': mode,
        'sources': sources,
        'history_count': len(chat_history)
    })


@app.route('/api/upload', methods=['POST'])
def upload_file():
    """上传文档到知识库"""
    if 'file' not in request.files:
        return jsonify({'error': '没有上传文件'}), 400
    file = request.files['file']
    if not file.filename:
        return jsonify({'error': '文件名为空'}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.md', '.pdf', '.docx']:
        return jsonify({'error': f'不支持的格式: {ext}，仅支持 txt/md/pdf/docx'}), 400

    # 保存文件
    filename = file.filename
    filepath = os.path.join(UPLOAD_FOLDER, filename)
    file.save(filepath)

    # 解析文档
    text = parse_file(filepath)
    if text.startswith('[PDF解析失败]') or text.startswith('[DOCX解析失败]') or text.startswith('[不支持的文件格式]'):
        os.remove(filepath)
        return jsonify({'error': text}), 400

    # 分块
    chunks = split_documents(text, filename)
    if not chunks:
        os.remove(filepath)
        return jsonify({'error': '文档解析结果为空，请检查文件内容'}), 400

    # 添加到知识库
    doc_record = {
        'id': str(uuid.uuid4())[:8],
        'source': filename,
        'chunks': chunks,
        'upload_time': datetime.now().strftime('%Y-%m-%d %H:%M:%S')
    }
    doc_list.append(doc_record)
    rebuild_index()

    return jsonify({
        'success': True,
        'message': f'文档"{filename}"上传成功，共解析 {len(chunks)} 个文本片段',
        'doc_id': doc_record['id'],
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks']) for d in doc_list)
    })


@app.route('/api/docs', methods=['GET'])
def get_docs():
    """获取知识库文档列表"""
    result = []
    for doc in doc_list:
        result.append({
            'id': doc['id'],
            'source': doc['source'],
            'chunk_count': len(doc['chunks']),
            'upload_time': doc.get('upload_time', '')
        })
    return jsonify({'docs': result, 'total': len(result)})


@app.route('/api/docs/<doc_id>', methods=['DELETE'])
def delete_doc(doc_id):
    """删除知识库中的文档"""
    global doc_list
    original_len = len(doc_list)
    doc_list = [d for d in doc_list if d['id'] != doc_id]
    if len(doc_list) == original_len:
        return jsonify({'error': '文档不存在'}), 404
    rebuild_index()
    return jsonify({
        'success': True,
        'message': '文档已删除',
        'total_docs': len(doc_list),
        'total_chunks': sum(len(d['chunks']) for d in doc_list)
    })


@app.route('/api/history', methods=['GET'])
def get_history():
    """获取对话历史"""
    return jsonify({'history': chat_history, 'total': len(chat_history)})


@app.route('/api/history/clear', methods=['POST'])
def clear_history():
    """清空对话历史"""
    global chat_history